"""DOCX document content extractor using python-docx."""

from pathlib import Path
from typing import Generator, Union

import docx

from src.ingestion.detectors.base import DetectionResult, FileType
from src.ingestion.extractors.base import BaseExtractor
from src.ingestion.extractors.models import ContentType, ExtractedElement, ExtractedStream, ExtractionStatus


class DOCXExtractor(BaseExtractor):
    """Extracts paragraphs, headings, and tables from DOCX files using python-docx."""

    def __init__(self):
        super().__init__(target_file_type=FileType.DOCX)

    def _extract_implementation(self, path: Path, detection_result: DetectionResult) -> ExtractedStream:
        try:
            doc = docx.Document(path)
            metadata = {
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables),
            }

            generator = self._element_generator(doc)
            return ExtractedStream(
                file_path=str(path),
                file_type=FileType.DOCX,
                element_generator=generator,
                metadata=metadata,
            )

        except Exception as e:
            return ExtractedStream(
                file_path=str(path),
                file_type=FileType.DOCX,
                element_generator=iter([]),
                status=ExtractionStatus.FAILED,
                error_message=f"Failed to open DOCX document: {str(e)}",
            )

    def _element_generator(self, doc: docx.Document) -> Generator[ExtractedElement, None, None]:
        element_idx = 0

        # Extract paragraphs & headings
        for p_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""
            content_type = ContentType.PARAGRAPH
            heading_level = None

            if style_name.startswith("Heading"):
                content_type = ContentType.HEADING
                try:
                    heading_level = int(style_name.replace("Heading", "").strip())
                except ValueError:
                    heading_level = 1

            yield ExtractedElement(
                element_index=element_idx,
                content_type=content_type,
                text_content=text,
                structural_context={
                    "paragraph_index": p_idx,
                    "style": style_name,
                    "heading_level": heading_level,
                },
            )
            element_idx += 1

        # Extract tables
        for t_idx, table in enumerate(doc.tables):
            table_rows = []
            for row in table.rows:
                row_cells = [cell.text.strip() for cell in row.cells]
                if any(row_cells):
                    table_rows.append(" | ".join(row_cells))

            if not table_rows:
                continue

            table_text = "\n".join(table_rows)
            yield ExtractedElement(
                element_index=element_idx,
                content_type=ContentType.TABLE,
                text_content=table_text,
                structural_context={
                    "table_index": t_idx,
                    "row_count": len(table.rows),
                    "col_count": len(table.columns),
                },
            )
            element_idx += 1
